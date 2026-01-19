"""
Service d'extraction de texte des documents PDF et Word
"""
from __future__ import annotations

from typing import Tuple, List
import io

# PDF: on privilégie PyMuPDF si dispo, sinon fallback pypdf (pur Python)
try:
    import fitz  # type: ignore  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None  # type: ignore

from pypdf import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

# Sur Windows, on peut pointer explicitement vers le binaire tesseract pour éviter
# les soucis de PATH. Adapte ce chemin si Tesseract est installé ailleurs.
try:  # pragma: no cover - dépend de l'environnement
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
except Exception as e:  # pragma: no cover
    print(f"Impossible de configurer pytesseract_cmd: {e}")


class TextExtractor:
    """Extracteur de texte pour PDF et Word"""
    
    @staticmethod
    def _ocr_image(image: Image.Image) -> str:
        """
        Effectue un OCR sur une image via Tesseract.
        On suppose que Tesseract est installé sur la machine (binaire système).
        """
        try:
            # Langue française prioritaire (à ajuster si besoin).
            text = pytesseract.image_to_string(image, lang="fra+eng")
            return text or ""
        except Exception as e:  # pragma: no cover - dépend du binaire système
            print(f"Erreur OCR Tesseract: {e}")
            return ""
    
    @staticmethod
    def extract_from_pdf(file_content: bytes) -> Tuple[str, bool]:
        """
        Extrait le texte d'un fichier PDF.
        
        Args:
            file_content: Contenu binaire du fichier PDF
            
        Returns:
            Tuple (texte extrait, succès)
        """
        # 1) PyMuPDF (meilleur rendu) si disponible
        if fitz is not None:
            try:
                doc = fitz.open(stream=file_content, filetype="pdf")
                text_parts: List[str] = []

                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    if text and text.strip():
                        # Texte natif disponible : on l'utilise tel quel
                        text_parts.append(text)
                    else:
                        # Pas de texte extrait → tentative d'OCR sur l'image de la page
                        try:
                            pix = page.get_pixmap()
                            img_bytes = pix.tobytes("png")
                            image = Image.open(io.BytesIO(img_bytes))
                            ocr_text = TextExtractor._ocr_image(image)
                            if ocr_text and ocr_text.strip():
                                text_parts.append(ocr_text)
                        except Exception as e:  # pragma: no cover - dépend du runtime
                            print(f"Erreur génération image pour OCR (PyMuPDF): {e}")

                doc.close()
                # Si on n'a vraiment rien récupéré, on indiquera un échec
                if not text_parts:
                    return "", False
                return "\n".join(text_parts), True
            except Exception as e:
                print(f"Erreur extraction PDF (PyMuPDF): {e}")

        # 2) Fallback pypdf (pur Python, plus compatible)
        try:
            reader = PdfReader(io.BytesIO(file_content))
            text_parts: List[str] = []
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    text_parts.append(text)

            # Avec pypdf on ne gère pas l'OCR directement (pas de rendu image ici).
            # Si aucun texte n'est trouvé, on signale un échec pour laisser la couche supérieure décider.
            if not text_parts:
                return "", False
            return "\n".join(text_parts), True
        except Exception as e:
            print(f"Erreur extraction PDF (pypdf): {e}")
            return "", False
    
    @staticmethod
    def extract_from_docx(file_content: bytes) -> Tuple[str, bool]:
        """
        Extrait le texte d'un fichier Word (.docx).
        
        Args:
            file_content: Contenu binaire du fichier Word
            
        Returns:
            Tuple (texte extrait, succès)
        """
        try:
            # Ouvrir le document Word depuis les bytes
            doc = DocxDocument(io.BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraire aussi le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            return full_text, True
            
        except Exception as e:
            print(f"Erreur extraction Word: {e}")
            return "", False
    
    @staticmethod
    def extract(file_content: bytes, filename: str) -> Tuple[str, bool]:
        """
        Extrait le texte d'un fichier selon son extension.
        
        Args:
            file_content: Contenu binaire du fichier
            filename: Nom du fichier (pour déterminer le type)
            
        Returns:
            Tuple (texte extrait, succès)
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith(".pdf"):
            return TextExtractor.extract_from_pdf(file_content)
        elif filename_lower.endswith(".docx"):
            return TextExtractor.extract_from_docx(file_content)
        elif filename_lower.endswith(".doc"):
            # Les fichiers .doc anciens ne sont pas supportés par python-docx
            # On retourne une chaîne vide mais on ne considère pas ça comme un échec
            return "", True
        else:
            return "", False

